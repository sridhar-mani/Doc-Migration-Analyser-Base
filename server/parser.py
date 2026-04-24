from docx import Document
import re
import fitz

def clean_uc(text:str ) -> str:
    cleaned_str = re.sub(r"[\uE000-\uF8FF]","",text)
    return re.sub(r"\s+", " ", cleaned_str).strip()



def calculate_base_metrics( text_blocks: list[str]):
    full_text = "\n".join(text_blocks)
    words = re.findall(r"\b\w+\b", full_text)
    word_cnt = len(words)
    para_cnt = len(text_blocks)

    avg_words_per_para = round(word_cnt / para_cnt) if para_cnt else 0

    return {
        "raw_text": full_text,
        "word_count": word_cnt,
        "paragraph_count": para_cnt,
        "avg_words_per_paragraph": avg_words_per_para,
    }


def collect_pdf_font_sizes(doc):
    return [
        round(span.get("size"))
        for page in doc
        for block in page.get_text("dict").get("blocks", [])
        if block.get("type") == 0
        for line in block.get("lines", [])
        for span in line.get("spans", [])
    ]


def block_is_heading(spans, base_size: int) -> bool:
    for span in spans:
        if span["size"] > base_size + 1.5 or "bold" in span["font"].lower():
            return True
    return False


def extract_pdf_text_blocks(doc, base_size: int):
    text_blocks = []
    heading_count = 0

    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue

            spans = [span for line in block.get("lines", []) for span in line.get("spans", [])]
            cleaned_text = " ".join(span.get("text", "") for span in spans).strip()
            if not cleaned_text:
                continue

            text_blocks.append(cleaned_text)
            if block_is_heading(spans, base_size) and len(cleaned_text.split()) < 15:
                heading_count += 1

    return text_blocks, heading_count

def parse_pdf(file_path: str):
    doc = fitz.open(file_path)
    font_sizes = collect_pdf_font_sizes(doc)
    base_size = max(set(font_sizes), key=font_sizes.count) if font_sizes else 12
    text_blocks, heading_count = extract_pdf_text_blocks(doc, base_size)

    metrics = calculate_base_metrics(text_blocks)
    metrics["heading_count"] = heading_count
    metrics["total_pages"] = len(doc)
    if metrics["word_count"] == 0:
        raise ValueError("Document is empty or contains no extractable text.")
    return metrics


def get_docx_page_count(doc: Document, word_count: int) -> int:
    page_count = 1

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            xml_str = run._element.xml
            if '<w:lastRenderedPageBreak' in xml_str or '<w:br w:type="page"' in xml_str:
                page_count += 1

    if page_count == 1 and word_count > 300:
        return max(1, word_count // 300)

    return page_count


def parse_docx(file_path: str):
    doc = Document(file_path)
    text_blocks = []
    heading_count = 0
    for para in doc.paragraphs:
        cleaned_text = para.text.strip()
        if cleaned_text:
            text_blocks.append(cleaned_text)
            if para.style.name.startswith("Heading") or para.style.name.startswith("Title"):
                heading_count += 1

    metrics = calculate_base_metrics(text_blocks)
    metrics["total_pages"] = get_docx_page_count(doc, metrics["word_count"])
    metrics["heading_count"] = heading_count
    if metrics["word_count"] == 0:
        raise ValueError("Document contains no extractable text.")
    return metrics


def parse(file_path: str, type: str):
    if type == "pdf":
        return parse_pdf(file_path)
    if type == "docx":
        return parse_docx(file_path)
    raise ValueError(f"Unsupported file type: {type}")